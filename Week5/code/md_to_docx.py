"""
md_to_docx.py — Week5 Day27（从 Week4/code/md_to_docx.py 原样复用，逻辑与它一致）
把 Markdown 报告转成 .docx（沿用 Week1/2 交付惯例：周报出 docx 版本）。
Convert a Markdown report to .docx.

支持 / Supports: H1-H4 标题、表格、有序/无序列表、引用块、代码块、行内 **加粗** 与 `代码`、
内嵌图片 ![alt](path)。中文字体设为微软雅黑。
Handles headings, tables, ordered/unordered lists, blockquotes, code blocks,
inline bold/code, embedded images; CJK font = Microsoft YaHei.

用法 / Usage（仓库根目录 / from repo root）:
    .venv-vlm/Scripts/python.exe Week5/code/md_to_docx.py Week5/deliverables/第5周_多模态实践报告.md
    .venv-vlm/Scripts/python.exe Week5/code/md_to_docx.py <in.md> [out.docx]
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[2]
CJK_FONT = "Microsoft YaHei"


def set_cjk(run) -> None:
    """让中文也用指定字体（python-docx 需单独设 east-asia 字形）。"""
    run.font.name = CJK_FONT
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)


def add_runs(paragraph, text: str) -> None:
    """解析行内 **加粗** 与 `代码`，逐段加 run。"""
    # 按 **bold** 和 `code` 切分
    for part in re.split(r"(\*\*.+?\*\*|`.+?`)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.color.rgb = RGBColor(0xC0, 0x30, 0x30)
        else:
            run = paragraph.add_run(part)
        set_cjk(run)


def is_table_sep(line: str) -> bool:
    """匹配 |---|---| 分隔行。"""
    return bool(re.match(r"^\s*\|?[\s:|-]+\|[\s:|-]*$", line)) and "-" in line


def split_row(line: str) -> list:
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def convert(md_path: Path, out_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    # 正文默认字体
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)

    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # 代码块 ```
        if stripped.startswith("```"):
            i += 1
            code = []
            while i < n and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1  # 跳过结尾 ```
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            continue

        # 表格：当前行含 | 且下一行是分隔行
        if "|" in line and i + 1 < n and is_table_sep(lines[i + 1]):
            header = split_row(line)
            i += 2  # 跳过表头与分隔行
            rows = []
            while i < n and "|" in lines[i] and lines[i].strip():
                rows.append(split_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Light Grid Accent 1"
            for j, h in enumerate(header):
                add_runs(table.rows[0].cells[j].paragraphs[0], h)
            for row in rows:
                cells = table.add_row().cells
                for j in range(len(header)):
                    add_runs(cells[j].paragraphs[0], row[j] if j < len(row) else "")
            doc.add_paragraph()
            continue

        # 标题
        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            h = doc.add_heading(level=min(level, 4))
            add_runs(h, m.group(2))
            i += 1
            continue

        # 内嵌图片 ![alt](path)
        img = re.match(r"^!\[.*?\]\((.+?)\)", stripped)
        if img:
            rel = img.group(1)
            ipath = (md_path.parent / rel)
            if ipath.exists():
                try:
                    from docx.shared import Inches
                    doc.add_picture(str(ipath), width=Inches(6))
                except Exception:
                    doc.add_paragraph(f"[图片: {rel}]")
            i += 1
            continue

        # 引用块 >
        if stripped.startswith(">"):
            p = doc.add_paragraph(style="Intense Quote")
            add_runs(p, stripped.lstrip("> ").strip())
            i += 1
            continue

        # 无序列表 - / *
        m = re.match(r"^[-*]\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, m.group(1))
            i += 1
            continue

        # 有序列表 1.
        m = re.match(r"^\d+\.\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, m.group(1))
            i += 1
            continue

        # 空行
        if not stripped:
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        add_runs(p, stripped)
        i += 1

    doc.save(str(out_path))
    print(f"[OK] {out_path.relative_to(ROOT)}")


def main() -> None:
    if len(sys.argv) < 2:
        raise SystemExit("用法: md_to_docx.py <in.md> [out.docx]")
    md_path = Path(sys.argv[1])
    if not md_path.is_absolute():
        md_path = ROOT / md_path
    out_path = Path(sys.argv[2]) if len(sys.argv) > 2 else md_path.with_suffix(".docx")
    if not out_path.is_absolute():
        out_path = ROOT / out_path
    convert(md_path, out_path)


if __name__ == "__main__":
    main()
